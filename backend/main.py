from fastapi import FastAPI, Depends, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime, timedelta
import os
import io
import json
import uuid

from . import models, schemas, crud
from .database import engine, get_db
from .logger import logger

# Create database tables
models.Base.metadata.create_all(bind=engine)

app = FastAPI(title="Work Logger API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ========== Project Endpoints ==========
@app.post("/api/projects/", response_model=schemas.Project)
def create_project(project: schemas.ProjectCreate, db: Session = Depends(get_db)):
    return crud.create_project(db, project)

@app.get("/api/projects/", response_model=List[schemas.Project])
def get_projects(year: Optional[int] = None, db: Session = Depends(get_db)):
    return crud.get_projects(db, year=year)

@app.get("/api/projects/{project_id}", response_model=schemas.Project)
def get_project(project_id: int, db: Session = Depends(get_db)):
    project = crud.get_project(db, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@app.delete("/api/projects/{project_id}")
def delete_project(project_id: int, db: Session = Depends(get_db)):
    success = crud.delete_project(db, project_id)
    if not success:
        raise HTTPException(status_code=404, detail="Project not found")
    return {"message": "Project deleted successfully"}

@app.patch("/api/projects/{project_id}", response_model=schemas.Project)
def update_project(
    project_id: int,
    project_update: schemas.ProjectUpdate,
    db: Session = Depends(get_db)
):
    project = crud.update_project(db, project_id, project_update)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

# ========== Annual Goal Endpoints ==========
@app.post("/api/annual-goals/", response_model=schemas.AnnualGoal)
def create_annual_goal(goal: schemas.AnnualGoalCreate, db: Session = Depends(get_db)):
    return crud.create_annual_goal(db, goal)

@app.get("/api/annual-goals/", response_model=List[schemas.AnnualGoal])
def get_annual_goals(
    project_id: Optional[int] = None,
    year: Optional[int] = None,
    db: Session = Depends(get_db)
):
    return crud.get_annual_goals(db, project_id, year)

@app.get("/api/annual-goals/{goal_id}", response_model=schemas.AnnualGoal)
def get_annual_goal(goal_id: int, db: Session = Depends(get_db)):
    goal = crud.get_annual_goal(db, goal_id)
    if not goal:
        raise HTTPException(status_code=404, detail="Annual goal not found")
    return goal

@app.delete("/api/annual-goals/{goal_id}")
def delete_annual_goal(goal_id: int, db: Session = Depends(get_db)):
    success = crud.delete_annual_goal(db, goal_id)
    if not success:
        raise HTTPException(status_code=404, detail="Annual goal not found")
    return {"message": "Annual goal deleted successfully"}

@app.patch("/api/annual-goals/{goal_id}", response_model=schemas.AnnualGoal)
def update_annual_goal(
    goal_id: int,
    goal_update: schemas.AnnualGoalUpdate,
    db: Session = Depends(get_db)
):
    goal = crud.update_annual_goal(db, goal_id, goal_update)
    if not goal:
        raise HTTPException(status_code=404, detail="Annual goal not found")
    return goal

# ========== Monthly Goal Endpoints ==========
@app.post("/api/monthly-goals/", response_model=schemas.MonthlyGoal)
def create_monthly_goal(goal: schemas.MonthlyGoalCreate, db: Session = Depends(get_db)):
    return crud.create_monthly_goal(db, goal)

@app.get("/api/monthly-goals/", response_model=List[schemas.MonthlyGoal])
def get_monthly_goals(
    annual_goal_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    return crud.get_monthly_goals(db, annual_goal_id)

@app.get("/api/monthly-goals/{goal_id}", response_model=schemas.MonthlyGoal)
def get_monthly_goal(goal_id: int, db: Session = Depends(get_db)):
    goal = crud.get_monthly_goal(db, goal_id)
    if not goal:
        raise HTTPException(status_code=404, detail="Monthly goal not found")
    return goal

@app.patch("/api/monthly-goals/{goal_id}/status")
def update_monthly_goal_status(
    goal_id: int,
    status: str = Query(..., regex="^(planned|in-progress|completed|terminated)$"),
    db: Session = Depends(get_db)
):
    goal = crud.update_monthly_goal_status(db, goal_id, status)
    if not goal:
        raise HTTPException(status_code=404, detail="Monthly goal not found")
    return goal

@app.patch("/api/monthly-goals/{goal_id}", response_model=schemas.MonthlyGoal)
def update_monthly_goal(
    goal_id: int,
    goal_update: schemas.MonthlyGoalUpdate,
    db: Session = Depends(get_db)
):
    try:
        goal = crud.update_monthly_goal(db, goal_id, goal_update)
        if not goal:
            raise HTTPException(status_code=404, detail="Monthly goal not found")
        return goal
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.delete("/api/monthly-goals/{goal_id}")
def delete_monthly_goal(goal_id: int, db: Session = Depends(get_db)):
    success = crud.delete_monthly_goal(db, goal_id)
    if not success:
        raise HTTPException(status_code=404, detail="Monthly goal not found")
    return {"message": "Monthly goal deleted successfully"}

# ========== Weekly Goal Endpoints ==========
@app.post("/api/weekly-goals/", response_model=schemas.WeeklyGoal)
def create_weekly_goal(goal: schemas.WeeklyGoalCreate, db: Session = Depends(get_db)):
    return crud.create_weekly_goal(db, goal)

@app.get("/api/weekly-goals/", response_model=List[schemas.WeeklyGoal])
def get_weekly_goals(db: Session = Depends(get_db)):
    return crud.get_weekly_goals(db)

@app.delete("/api/weekly-goals/{goal_id}")
def delete_weekly_goal(goal_id: int, db: Session = Depends(get_db)):
    success = crud.delete_weekly_goal(db, goal_id)
    if not success:
        raise HTTPException(status_code=404, detail="Weekly goal not found")
    return {"message": "Weekly goal deleted successfully"}

# ========== Task Endpoints ==========
@app.post("/api/tasks/", response_model=schemas.Task)
def create_task(task: schemas.TaskCreate, db: Session = Depends(get_db)):
    return crud.create_task(db, task)

@app.get("/api/tasks/", response_model=List[schemas.Task])
def get_tasks(
    date: Optional[str] = None,
    completed: Optional[bool] = None,
    db: Session = Depends(get_db)
):
    date_obj = None
    if date:
        try:
            date_obj = datetime.strptime(date, '%Y-%m-%d')
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    return crud.get_tasks(db, date_obj, completed)

@app.get("/api/tasks/today", response_model=List[schemas.Task])
def get_today_tasks(db: Session = Depends(get_db)):
    return crud.get_tasks(db, datetime.now())

@app.get("/api/tasks/search")
def search_tasks(
    q: str = Query(..., min_length=1),
    page: int = Query(default=1, ge=1),
    limit: int = Query(default=5, ge=1, le=100),
    db: Session = Depends(get_db)
):
    """Search tasks across all time ranges"""
    results = crud.search_tasks(db, query=q, page=page, limit=limit)
    return results

@app.get("/api/tasks/{task_id}", response_model=schemas.Task)
def get_task(task_id: int, db: Session = Depends(get_db)):
    task = crud.get_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return task

@app.patch("/api/tasks/{task_id}", response_model=schemas.Task)
def update_task(task_id: int, task_update: schemas.TaskUpdate, db: Session = Depends(get_db)):
    task = crud.update_task(db, task_id, task_update)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return task

@app.delete("/api/tasks/{task_id}")
def delete_task(task_id: int, db: Session = Depends(get_db)):
    success = crud.delete_task(db, task_id)
    if not success:
        raise HTTPException(status_code=404, detail="Task not found")
    return {"message": "Task deleted successfully"}

# ========== Experiment Endpoints ==========
@app.post("/api/experiments/", response_model=schemas.Experiment)
def create_experiment(experiment: schemas.ExperimentCreate, db: Session = Depends(get_db)):
    return crud.create_experiment(db, experiment)

@app.get("/api/experiments/", response_model=List[schemas.Experiment])
def get_experiments(status: Optional[str] = None, db: Session = Depends(get_db)):
    return crud.get_experiments(db, status)

@app.get("/api/experiments/{experiment_id}", response_model=schemas.Experiment)
def get_experiment(experiment_id: int, db: Session = Depends(get_db)):
    experiment = crud.get_experiment(db, experiment_id)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")
    return experiment

@app.patch("/api/experiments/{experiment_id}", response_model=schemas.Experiment)
def update_experiment(
    experiment_id: int,
    experiment_update: schemas.ExperimentUpdate,
    db: Session = Depends(get_db)
):
    experiment = crud.update_experiment(db, experiment_id, experiment_update)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")
    return experiment

@app.get("/api/experiments/{experiment_id}/tasks", response_model=List[schemas.Task])
def get_experiment_tasks(experiment_id: int, db: Session = Depends(get_db)):
    return crud.get_experiment_tasks(db, experiment_id)

@app.delete("/api/experiments/{experiment_id}")
def delete_experiment(experiment_id: int, db: Session = Depends(get_db)):
    success = crud.delete_experiment(db, experiment_id)
    if not success:
        raise HTTPException(status_code=404, detail="Experiment not found")
    return {"message": "Experiment deleted successfully"}

@app.post("/api/experiments/{experiment_id}/duplicate", response_model=schemas.Experiment)
def duplicate_experiment(experiment_id: int, db: Session = Depends(get_db)):
    """Duplicate an experiment with copied fields and cleared relationships"""
    experiment = crud.duplicate_experiment(db, experiment_id)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")
    return experiment

# ========== Experiment Image Endpoints ==========
@app.post("/api/experiments/{experiment_id}/images")
def upload_experiment_image(experiment_id: int, image_data: dict, db: Session = Depends(get_db)):
    """Upload an image to an experiment (base64 encoded)"""
    experiment = crud.get_experiment(db, experiment_id)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")

    # Parse existing images
    images = json.loads(experiment.images_json or '[]')

    # Create new image entry
    new_image = {
        'id': str(uuid.uuid4()),
        'filename': image_data.get('filename', 'image.png'),
        'data': image_data.get('base64_data', '')
    }
    images.append(new_image)

    # Update experiment
    experiment.images_json = json.dumps(images)
    db.commit()
    db.refresh(experiment)

    return new_image

@app.delete("/api/experiments/{experiment_id}/images/{image_id}")
def delete_experiment_image(experiment_id: int, image_id: str, db: Session = Depends(get_db)):
    """Delete an image from an experiment"""
    experiment = crud.get_experiment(db, experiment_id)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")

    # Parse existing images
    images = json.loads(experiment.images_json or '[]')

    # Remove the image with matching id
    images = [img for img in images if img.get('id') != image_id]

    # Update experiment
    experiment.images_json = json.dumps(images)
    db.commit()

    return {"message": "Image deleted successfully"}

@app.get("/api/experiments/{experiment_id}/images")
def get_experiment_images(experiment_id: int, db: Session = Depends(get_db)):
    """Get all images for an experiment"""
    experiment = crud.get_experiment(db, experiment_id)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")

    images = json.loads(experiment.images_json or '[]')
    return {"images": images}

# ========== Custom Tag Endpoints ==========
@app.get("/api/tags/", response_model=List[schemas.CustomTag])
def get_custom_tags(db: Session = Depends(get_db)):
    return crud.get_custom_tags(db)

# ========== Report Endpoints ==========
@app.get("/api/reports/daily")
def get_daily_report(date: str = Query(...), db: Session = Depends(get_db)):
    try:
        date_obj = datetime.strptime(date, '%Y-%m-%d')
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    return crud.get_daily_report(db, date_obj)

@app.get("/api/reports/weekly")
def get_weekly_report(week_start: str = Query(...), db: Session = Depends(get_db)):
    try:
        week_start_obj = datetime.strptime(week_start, '%Y-%m-%d')
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    return crud.get_weekly_report(db, week_start_obj)

@app.get("/api/reports/gantt")
def get_gantt_data(year: int = Query(default=datetime.now().year), db: Session = Depends(get_db)):
    return crud.get_gantt_data(db, year)

# ========== Export Endpoints ==========
@app.get("/api/export/excel/weekly")
def export_weekly_excel(week_start: str = Query(...), db: Session = Depends(get_db)):
    """Export weekly report to Excel"""
    try:
        week_start_obj = datetime.strptime(week_start, '%Y-%m-%d')
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")

    report_data = crud.get_weekly_report(db, week_start_obj)
    excel_file = generate_weekly_excel(report_data, db)

    return StreamingResponse(
        excel_file,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=weekly_report_{week_start}.xlsx"}
    )

@app.get("/api/export/excel/monthly")
def export_monthly_excel(month: int = Query(...), year: int = Query(...), db: Session = Depends(get_db)):
    """Export monthly report to Excel"""
    excel_file = generate_monthly_excel(month, year, db)

    return StreamingResponse(
        excel_file,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=monthly_report_{year}_{month:02d}.xlsx"}
    )

@app.get("/api/export/excel/annual")
def export_annual_excel(year: int = Query(...), db: Session = Depends(get_db)):
    """Export annual report to Excel"""
    gantt_data = crud.get_gantt_data(db, year)
    excel_file = generate_annual_excel(gantt_data, year)

    return StreamingResponse(
        excel_file,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=annual_report_{year}.xlsx"}
    )

def generate_weekly_excel(report_data, db: Session):
    """Generate Excel file for weekly report"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Weekly Report"

    # Styles
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
    title_font = Font(bold=True, size=14)
    monthly_goal_font = Font(bold=True, size=11, color="2C5282")
    monthly_goal_fill = PatternFill(start_color="BEE3F8", end_color="BEE3F8", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Date range in first row
    ws['A1'] = f"Weekly Report: {report_data['week_start']} to {report_data['week_end']}"
    ws['A1'].font = title_font
    ws.merge_cells('A1:D1')

    # Column headers: Monthly Goal | Weekly Goal | Task Description | Status
    headers = ['Monthly Goal', 'Weekly Goal', 'Task Description', 'Status']
    row = 2
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center')

    # Group tasks by monthly goal, then by weekly goal
    monthly_goal_groups = {}
    tasks = report_data.get('tasks', [])

    for task in tasks:
        monthly_goal_id = task.monthly_goal_id or 0
        weekly_goal_id = task.weekly_goal_id or 0

        if monthly_goal_id not in monthly_goal_groups:
            mg_name = ''
            if task.monthly_goal_id:
                mg = db.query(models.MonthlyGoal).filter(models.MonthlyGoal.id == task.monthly_goal_id).first()
                mg_name = mg.name if mg else ''
            else:
                mg_name = 'Unassigned'
            monthly_goal_groups[monthly_goal_id] = {
                'name': mg_name,
                'weekly_goals': {}
            }

        if weekly_goal_id not in monthly_goal_groups[monthly_goal_id]['weekly_goals']:
            wg_name = ''
            if task.weekly_goal_id:
                wg = db.query(models.WeeklyGoal).filter(models.WeeklyGoal.id == task.weekly_goal_id).first()
                wg_name = wg.name if wg else ''
            else:
                wg_name = 'Unassigned'
            monthly_goal_groups[monthly_goal_id]['weekly_goals'][weekly_goal_id] = {
                'name': wg_name,
                'tasks': []
            }

        monthly_goal_groups[monthly_goal_id]['weekly_goals'][weekly_goal_id]['tasks'].append(task)

    # Write grouped data
    for mg_id, mg_data in monthly_goal_groups.items():
        first_mg_row = True
        mg_start_row = row + 1

        for wg_id, wg_data in mg_data['weekly_goals'].items():
            first_wg_row = True

            for task in wg_data['tasks']:
                row += 1

                # Monthly goal name (only on first row of monthly goal group)
                if first_mg_row:
                    cell = ws.cell(row=row, column=1, value=mg_data['name'])
                    cell.border = border
                    cell.font = monthly_goal_font
                    cell.fill = monthly_goal_fill
                    first_mg_row = False
                else:
                    ws.cell(row=row, column=1, value='').border = border

                # Weekly goal name (only on first row of weekly goal group)
                if first_wg_row:
                    ws.cell(row=row, column=2, value=wg_data['name']).border = border
                    first_wg_row = False
                else:
                    ws.cell(row=row, column=2, value='').border = border

                # Task description and status
                ws.cell(row=row, column=3, value=task.title or task.description or '').border = border
                ws.cell(row=row, column=4, value='✓ Done' if task.completed else 'Pending').border = border

        # Merge monthly goal cells
        if row >= mg_start_row:
            ws.merge_cells(f'A{mg_start_row}:A{row}')

    # Auto-adjust column widths
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 45
    ws.column_dimensions['D'].width = 12

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_monthly_excel(month, year, db: Session):
    """Generate Excel file for monthly report"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Monthly Report"

    # Styles
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    month_names = ['', 'January', 'February', 'March', 'April', 'May', 'June',
                   'July', 'August', 'September', 'October', 'November', 'December']

    # Month in first row
    ws['A1'] = f"Monthly Report: {month_names[month]} {year}"
    ws['A1'].font = title_font
    ws.merge_cells('A1:C1')

    # Column headers: Annual Goal | Monthly Goals | Status
    headers = ['Annual Goal', 'Monthly Goals', 'Status']
    row = 2
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center')

    # Get month range for filtering
    month_start = datetime(year, month, 1)
    if month == 12:
        month_end = datetime(year + 1, 1, 1) - timedelta(days=1)
    else:
        month_end = datetime(year, month + 1, 1) - timedelta(days=1)

    # Get all monthly goals that overlap with this month
    all_monthly_goals = db.query(models.MonthlyGoal).all()
    relevant_goals = [mg for mg in all_monthly_goals
                     if not (mg.end_date < month_start.date() or mg.start_date > month_end.date())]

    # Group by annual goal
    annual_goal_map = {}
    for mg in relevant_goals:
        ag = db.query(models.AnnualGoal).filter(models.AnnualGoal.id == mg.annual_goal_id).first()
        if ag:
            if ag.id not in annual_goal_map:
                annual_goal_map[ag.id] = {'name': ag.name, 'monthly_goals': []}
            annual_goal_map[ag.id]['monthly_goals'].append({
                'name': mg.name,
                'status': mg.status
            })

    # Styles for annual goal
    annual_goal_font = Font(bold=True, size=11, color="2C5282")
    annual_goal_fill = PatternFill(start_color="BEE3F8", end_color="BEE3F8", fill_type="solid")

    # Write data rows with grouped annual goals
    for ag_id, ag_data in annual_goal_map.items():
        ag_start_row = row + 1
        first_row = True

        for mg in ag_data['monthly_goals']:
            row += 1

            # Annual goal name (only on first row)
            if first_row:
                cell = ws.cell(row=row, column=1, value=ag_data['name'])
                cell.border = border
                cell.font = annual_goal_font
                cell.fill = annual_goal_fill
                first_row = False
            else:
                ws.cell(row=row, column=1, value='').border = border

            ws.cell(row=row, column=2, value=mg['name']).border = border
            ws.cell(row=row, column=3, value=mg['status']).border = border

        # Merge annual goal cells
        if row >= ag_start_row:
            ws.merge_cells(f'A{ag_start_row}:A{row}')

    # Auto-adjust column widths
    ws.column_dimensions['A'].width = 35
    ws.column_dimensions['B'].width = 35
    ws.column_dimensions['C'].width = 15

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_annual_excel(gantt_data, year):
    """Generate Excel file for annual report"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Annual Report"

    # Styles
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Title
    ws['A1'] = f"Annual Report: {year}"
    ws['A1'].font = title_font
    ws.merge_cells('A1:F1')

    # Headers
    headers = ['Project', 'Annual Goal', 'Monthly Goals', 'Total Tasks', 'Completed', 'Progress']
    row = 3
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border

    # Data
    for project in gantt_data:
        for goal in project.get('annual_goals', []):
            row += 1
            ws.cell(row=row, column=1, value=project['name']).border = border
            ws.cell(row=row, column=2, value=goal['name']).border = border
            ws.cell(row=row, column=3, value=len(goal.get('monthly_goals', []))).border = border
            ws.cell(row=row, column=4, value=goal.get('total_tasks', 0)).border = border
            ws.cell(row=row, column=5, value=goal.get('completed_tasks', 0)).border = border

            total = goal.get('total_tasks', 0)
            completed = goal.get('completed_tasks', 0)
            progress = f"{(completed/total*100):.1f}%" if total > 0 else "0%"
            ws.cell(row=row, column=6, value=progress).border = border

    # Auto-adjust column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 12

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# ========== PDF Export Endpoints ==========
@app.get("/api/export/pdf/weekly")
def export_weekly_pdf(week_start: str = Query(...), db: Session = Depends(get_db)):
    """Export weekly report to PDF"""
    try:
        week_start_obj = datetime.strptime(week_start, '%Y-%m-%d')
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")

    report_data = crud.get_weekly_report(db, week_start_obj)
    pdf_file = generate_weekly_pdf(report_data)

    return StreamingResponse(
        pdf_file,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=weekly_report_{week_start}.pdf"}
    )

@app.get("/api/export/pdf/monthly")
def export_monthly_pdf(month: int = Query(...), year: int = Query(...), db: Session = Depends(get_db)):
    """Export monthly report to PDF"""
    month_start = datetime(year, month, 1)
    if month == 12:
        month_end = datetime(year + 1, 1, 1) - timedelta(days=1)
    else:
        month_end = datetime(year, month + 1, 1) - timedelta(days=1)

    tasks = crud.get_tasks(db)
    monthly_tasks = [t for t in tasks if t.date and month_start <= t.date <= month_end]
    monthly_goals = crud.get_monthly_goals(db)

    pdf_file = generate_monthly_pdf(monthly_tasks, monthly_goals, month, year)

    return StreamingResponse(
        pdf_file,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=monthly_report_{year}_{month:02d}.pdf"}
    )

@app.get("/api/export/pdf/annual")
def export_annual_pdf(year: int = Query(...), db: Session = Depends(get_db)):
    """Export annual report to PDF"""
    gantt_data = crud.get_gantt_data(db, year)
    pdf_file = generate_annual_pdf(gantt_data, year)

    return StreamingResponse(
        pdf_file,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=annual_report_{year}.pdf"}
    )

@app.get("/api/export/word/daily")
def export_daily_word(date: str = Query(...), db: Session = Depends(get_db)):
    """Export daily report to Word document"""
    try:
        date_obj = datetime.strptime(date, '%Y-%m-%d')
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")

    word_file = generate_daily_word(date_obj, db)

    return StreamingResponse(
        word_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=daily_report_{date}.docx"}
    )

def generate_weekly_pdf(report_data):
    """Generate PDF file for weekly report"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#667eea'),
        spaceAfter=20
    )
    elements.append(Paragraph(
        f"Weekly Report: {report_data['week_start']} to {report_data['week_end']}",
        title_style
    ))

    # Summary
    elements.append(Paragraph("Summary", styles['Heading2']))
    summary_data = [
        ['Metric', 'Value'],
        ['Total Tasks', str(report_data['total_tasks'])],
        ['Completed', str(report_data['completed_tasks'])],
        ['Completion Rate', f"{report_data['completion_rate']:.1f}%"]
    ]
    summary_table = Table(summary_data, colWidths=[200, 100])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f7fafc')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 20))

    # Tasks table
    elements.append(Paragraph("Tasks", styles['Heading2']))
    task_data = [['Title', 'Date', 'Status']]
    for task in report_data.get('tasks', []):
        title = task.title[:40] + '...' if len(task.title) > 40 else task.title
        task_data.append([
            title,
            task.date.strftime('%Y-%m-%d') if task.date else '',
            'Completed' if task.completed else 'Pending'
        ])

    if len(task_data) > 1:
        task_table = Table(task_data, colWidths=[300, 100, 80])
        task_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(task_table)
    else:
        elements.append(Paragraph("No tasks for this week.", styles['Normal']))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_monthly_pdf(tasks, monthly_goals, month, year):
    """Generate PDF file for monthly report"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    month_names = ['', 'January', 'February', 'March', 'April', 'May', 'June',
                   'July', 'August', 'September', 'October', 'November', 'December']

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#667eea'),
        spaceAfter=20
    )
    elements.append(Paragraph(f"Monthly Report: {month_names[month]} {year}", title_style))

    # Summary
    total_tasks = len(tasks)
    completed_tasks = sum(1 for t in tasks if t.completed)
    completion_rate = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0

    elements.append(Paragraph("Summary", styles['Heading2']))
    summary_data = [
        ['Metric', 'Value'],
        ['Total Tasks', str(total_tasks)],
        ['Completed', str(completed_tasks)],
        ['Completion Rate', f"{completion_rate:.1f}%"]
    ]
    summary_table = Table(summary_data, colWidths=[200, 100])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f7fafc')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 20))

    # Tasks table
    elements.append(Paragraph("Tasks", styles['Heading2']))
    task_data = [['Title', 'Date', 'Status']]
    for task in tasks:
        title = task.title[:35] + '...' if len(task.title) > 35 else task.title
        task_data.append([
            title,
            task.date.strftime('%Y-%m-%d') if task.date else '',
            'Completed' if task.completed else 'Pending'
        ])

    if len(task_data) > 1:
        task_table = Table(task_data, colWidths=[300, 100, 80])
        task_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(task_table)
    else:
        elements.append(Paragraph("No tasks for this month.", styles['Normal']))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_annual_pdf(gantt_data, year):
    """Generate PDF file for annual report"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#667eea'),
        spaceAfter=20
    )
    elements.append(Paragraph(f"Annual Report: {year}", title_style))

    # Goals table
    elements.append(Paragraph("Goals Summary", styles['Heading2']))
    goal_data = [['Project', 'Annual Goal', 'Tasks', 'Completed', 'Progress']]

    for project in gantt_data:
        for goal in project.get('annual_goals', []):
            total = goal.get('total_tasks', 0)
            completed = goal.get('completed_tasks', 0)
            progress = f"{(completed/total*100):.0f}%" if total > 0 else "0%"
            goal_data.append([
                project['name'][:20],
                goal['name'][:25],
                str(total),
                str(completed),
                progress
            ])

    if len(goal_data) > 1:
        goal_table = Table(goal_data, colWidths=[100, 150, 60, 70, 60])
        goal_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (2, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(goal_table)
    else:
        elements.append(Paragraph("No goals data for this year.", styles['Normal']))

    doc.build(elements)
    buffer.seek(0)
    return buffer

# ========== Logging Endpoints ==========
@app.post("/api/logs/")
def write_logs(logs_data: dict):
    """Receive logs from frontend and write them to log files."""
    logs = logs_data.get("logs", [])
    for log_entry in logs:
        level = log_entry.get("level", "INFO")
        message = log_entry.get("message", "")
        source = log_entry.get("source", "frontend")
        data = log_entry.get("data")

        logger.write_log(level, message, source, data)

    return {"status": "ok", "count": len(logs)}

@app.get("/api/logs/")
def get_logs(
    date: Optional[str] = None,
    level: Optional[str] = None,
    source: Optional[str] = None,
    limit: int = Query(default=100, le=1000),
    offset: int = Query(default=0, ge=0)
):
    """Retrieve logs from files."""
    logs = logger.get_logs(date=date, level=level, source=source, limit=limit, offset=offset)
    return {"logs": logs, "count": len(logs)}

@app.get("/api/logs/files")
def get_log_files():
    """Get list of available log files."""
    files = logger.get_log_files()
    return {"files": files, "count": len(files)}

@app.delete("/api/logs/")
def clear_logs(date: Optional[str] = None):
    """Clear logs for a specific date or all logs."""
    success = logger.clear_logs(date)
    if success:
        return {"status": "ok", "message": "Logs cleared successfully"}
    raise HTTPException(status_code=500, detail="Failed to clear logs")

# ========== Health Check ==========
@app.get("/api/health")
def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

def generate_daily_word(date_obj: datetime, db: Session):
    """Generate Word file for daily report"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import re
    import html

    # Helper functions for parsing HTML and JSON
    def html_to_text(html_content):
        """Convert HTML to plain text, preserving paragraphs"""
        if not html_content:
            return ""

        # Replace block-level tags with newlines
        html_content = re.sub(r'</p>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<p[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<br/?>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<div[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</div>', '\n', html_content, flags=re.IGNORECASE)

        # Strip formatting tags but keep content
        html_content = re.sub(r'<[^>]+>', '', html_content)

        # Decode HTML entities
        html_content = html.unescape(html_content)

        # Clean up multiple newlines
        html_content = re.sub(r'\n\s*\n', '\n\n', html_content)

        return html_content.strip()

    def parse_experiment_field(field_content):
        """Parse field - can be plain HTML or JSON with text and tables"""
        if not field_content:
            return None, []

        try:
            # Try to parse as JSON
            data = json.loads(field_content)
            text = data.get('text', '')
            tables = data.get('tables', [])
            return text, tables
        except (json.JSONDecodeError, ValueError):
            # If not JSON, treat as plain HTML
            return field_content, []

    def add_html_paragraph(doc, html_content, style=None):
        """Add HTML content as formatted paragraph(s)"""
        if not html_content:
            return

        text_content = html_to_text(html_content)
        if text_content:
            for para_text in text_content.split('\n\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip(), style=style)

    def add_table_to_doc(doc, table_data):
        """Add a table to the Word document"""
        if not table_data or 'data' not in table_data:
            return

        headers = table_data['data'].get('headers', [])
        rows = table_data['data'].get('rows', [])

        if not headers or not rows:
            return

        # Add table title if present
        if 'title' in table_data and table_data['title']:
            title_para = doc.add_paragraph(table_data['title'])
            title_para.runs[0].font.bold = True
            title_para.runs[0].font.size = Pt(11)

        # Create table
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            if idx < len(header_cells):
                header_cells[idx].text = str(header)
                for paragraph in header_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = str(cell_value)

        doc.add_paragraph()  # Spacing after table

    doc = Document()

    # Title
    title = doc.add_heading(f'Daily Report - {date_obj.strftime("%Y-%m-%d")}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Get tasks for the date
    tasks = crud.get_tasks(db, date=date_obj)

    if not tasks:
        doc.add_paragraph('No tasks for this date.')
    else:
        # Group tasks by weekly goal
        tasks_by_weekly_goal = {}
        unassigned_tasks = []

        for task in tasks:
            if task.weekly_goal_id:
                if task.weekly_goal_id not in tasks_by_weekly_goal:
                    tasks_by_weekly_goal[task.weekly_goal_id] = []
                tasks_by_weekly_goal[task.weekly_goal_id].append(task)
            else:
                unassigned_tasks.append(task)

        # Render tasks by weekly goal
        for weekly_goal_id, goal_tasks in tasks_by_weekly_goal.items():
            weekly_goal = crud.get_weekly_goal(db, weekly_goal_id)
            if weekly_goal:
                # Add weekly goal as Heading 2
                doc.add_heading(weekly_goal.name, 2)

                for task in goal_tasks:
                    # Add task title as Heading 3
                    doc.add_heading(task.title, 3)

                    # Add task description
                    if task.description:
                        doc.add_paragraph(task.description)

                    # Add linked experiments
                    if task.experiments:
                        for exp in task.experiments:
                            doc.add_paragraph(f'Linked Experiment: {exp.name}', style='Intense Quote')

                            # Add experiment sections
                            # Process Hypothesis (plain HTML)
                            if exp.hypothesis:
                                doc.add_heading('Hypothesis', 4)
                                add_html_paragraph(doc, exp.hypothesis)

                            # Process Methodology (JSON with text + tables)
                            if exp.methodology:
                                doc.add_heading('Methodology', 4)
                                method_text, method_tables = parse_experiment_field(exp.methodology)
                                if method_text:
                                    add_html_paragraph(doc, method_text)
                                for table_data in method_tables:
                                    add_table_to_doc(doc, table_data)

                            # Process Results (JSON with text + tables)
                            if exp.results:
                                doc.add_heading('Results', 4)
                                results_text, results_tables = parse_experiment_field(exp.results)
                                if results_text:
                                    add_html_paragraph(doc, results_text)
                                for table_data in results_tables:
                                    add_table_to_doc(doc, table_data)

                            # Process Conclusion (plain HTML)
                            if exp.conclusion:
                                doc.add_heading('Conclusion', 4)
                                add_html_paragraph(doc, exp.conclusion)

                            # Process Progress Notes (plain HTML)
                            if exp.progress_notes:
                                doc.add_heading('Progress Notes', 4)
                                add_html_paragraph(doc, exp.progress_notes)

                    doc.add_paragraph()  # Add spacing

        # Render unassigned tasks
        if unassigned_tasks:
            doc.add_heading('Unassigned Tasks', 2)

            for task in unassigned_tasks:
                doc.add_heading(task.title, 3)

                if task.description:
                    doc.add_paragraph(task.description)

                if task.experiments:
                    for exp in task.experiments:
                        doc.add_paragraph(f'Linked Experiment: {exp.name}', style='Intense Quote')

                        # Process Hypothesis (plain HTML)
                        if exp.hypothesis:
                            doc.add_heading('Hypothesis', 4)
                            add_html_paragraph(doc, exp.hypothesis)

                        # Process Methodology (JSON with text + tables)
                        if exp.methodology:
                            doc.add_heading('Methodology', 4)
                            method_text, method_tables = parse_experiment_field(exp.methodology)
                            if method_text:
                                add_html_paragraph(doc, method_text)
                            for table_data in method_tables:
                                add_table_to_doc(doc, table_data)

                        # Process Results (JSON with text + tables)
                        if exp.results:
                            doc.add_heading('Results', 4)
                            results_text, results_tables = parse_experiment_field(exp.results)
                            if results_text:
                                add_html_paragraph(doc, results_text)
                            for table_data in results_tables:
                                add_table_to_doc(doc, table_data)

                        # Process Conclusion (plain HTML)
                        if exp.conclusion:
                            doc.add_heading('Conclusion', 4)
                            add_html_paragraph(doc, exp.conclusion)

                        # Process Progress Notes (plain HTML)
                        if exp.progress_notes:
                            doc.add_heading('Progress Notes', 4)
                            add_html_paragraph(doc, exp.progress_notes)

                doc.add_paragraph()

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

# ========== Frontend Routes ==========
# Serve frontend files
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")

# Mount static files
app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")

@app.get("/")
def serve_index():
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))

@app.get("/api.js")
def serve_api_js():
    return FileResponse(os.path.join(FRONTEND_DIR, "api.js"))

@app.get("/{page}.html")
def serve_page(page: str):
    file_path = os.path.join(FRONTEND_DIR, f"{page}.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    raise HTTPException(status_code=404, detail="Page not found")
